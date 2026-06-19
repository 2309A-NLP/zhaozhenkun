"""
文档分类标签模块
Document Classification Tags Module
"""

import os
import time
from typing import Dict, List, Any, Optional
from .base import ConfigManager, logger


class DocumentClassifier:
    """文档分类器"""
    
    def __init__(self, config: ConfigManager):
        self.config = config
        self.tags_config = config.get('classification.tags', {})
        
        # 内容类型标签
        self.content_type_tags = self.tags_config.get('content_type', [
            '技术文档', '用户手册', 'API文档', '设计文档', 
            '会议纪要', '研究报告', '培训材料', '其他'
        ])
        
        # 质量标签
        self.quality_tags = self.tags_config.get('quality', [
            '高质量', '中等质量', '低质量', '待评估'
        ])
        
        # 处理状态标签
        self.processing_status_tags = self.tags_config.get('processing_status', [
            '待处理', '处理中', '已完成', '有问题'
        ])
    
    def classify_document(self, file_path: str, content: str = None) -> Dict[str, Any]:
        """
        对文档进行分类
        
        Args:
            file_path: 文件路径
            content: 文件内容（可选，如果不提供则自动读取）
            
        Returns:
            分类结果
        """
        try:
            # 如果没有提供内容，则读取文件
            if content is None:
                content = self._read_file_content(file_path)
            
            if not content:
                return self._get_default_classification(file_path, "无法读取文件内容")
            
            # 分析内容类型
            content_type = self._analyze_content_type(content, file_path)
            
            # 分析质量等级
            quality = self._analyze_quality(content, file_path)
            
            # 分析处理状态
            processing_status = self._analyze_processing_status(file_path)
            
            # 提取关键词
            keywords = self._extract_keywords(content)
            
            classification = {
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'content_type': content_type,
                'quality': quality,
                'processing_status': processing_status,
                'keywords': keywords,
                'classification_confidence': self._calculate_confidence(content, content_type),
                'classification_reason': self._get_classification_reason(content_type, quality)
            }
            
            logger.info(f"文档分类完成: {file_path} -> {content_type}/{quality}")
            return classification
            
        except Exception as e:
            logger.error(f"文档分类失败: {file_path}, 错误: {e}")
            return self._get_default_classification(file_path, str(e))
    
    def _analyze_content_type(self, content: str, file_path: str) -> str:
        """分析内容类型"""
        content_lower = content.lower()
        file_name = os.path.basename(file_path).lower()
        
        # 基于文件名和内容的关键词匹配
        type_indicators = {
            'API文档': ['api', '接口', 'endpoint', 'request', 'response', 'rest', 'graphql'],
            '用户手册': ['手册', '指南', '教程', '使用说明', '操作步骤', 'manual', 'guide'],
            '技术文档': ['技术', '架构', '设计', '实现', '代码', '算法', 'technical'],
            '设计文档': ['设计', '方案', '架构图', '流程图', 'design', 'specification'],
            '会议纪要': ['会议', '纪要', '讨论', '决议', 'meeting', 'minutes'],
            '研究报告': ['报告', '研究', '分析', '数据', 'report', 'analysis'],
            '培训材料': ['培训', '教程', '学习', '课程', 'training', 'tutorial']
        }
        
        # 统计各类型的匹配分数
        type_scores = {}
        for doc_type, indicators in type_indicators.items():
            score = 0
            for indicator in indicators:
                if indicator in content_lower or indicator in file_name:
                    score += 1
            type_scores[doc_type] = score
        
        # 返回得分最高的类型
        if type_scores:
            best_type = max(type_scores.items(), key=lambda x: x[1])
            if best_type[1] > 0:
                return best_type[0]
        
        return '其他'
    
    def _analyze_quality(self, content: str, file_path: str) -> str:
        """分析质量等级"""
        # 基于多个指标综合评估
        quality_score = 0
        
        # 1. 文件大小评估
        file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
        if file_size > 10000:  # 大于10KB
            quality_score += 1
        if file_size > 100000:  # 大于100KB
            quality_score += 1
        
        # 2. 内容长度评估
        content_length = len(content)
        if content_length > 1000:  # 大于1000字符
            quality_score += 1
        if content_length > 10000:  # 大于10000字符
            quality_score += 1
        
        # 3. 结构化程度评估
        lines = content.split('\n')
        non_empty_lines = [line.strip() for line in lines if line.strip()]
        if len(non_empty_lines) > 10:
            quality_score += 1
        
        # 4. 标题和段落结构
        has_headers = any(line.strip().startswith('#') for line in lines[:20])
        if has_headers:
            quality_score += 1
        
        # 5. 特殊字符和格式
        has_formatting = any(marker in content for marker in ['**', '__', '```', '- ', '* ', '1.'])
        if has_formatting:
            quality_score += 1
        
        # 根据分数判断质量等级
        if quality_score >= 5:
            return '高质量'
        elif quality_score >= 3:
            return '中等质量'
        elif quality_score >= 1:
            return '低质量'
        else:
            return '待评估'
    
    def _analyze_processing_status(self, file_path: str) -> str:
        """分析处理状态 — 基于文件修改时间和存在性判断"""
        try:
            stat = os.stat(file_path)
            mtime = stat.st_mtime
            now = time.time()
            days_since_modified = (now - mtime) / 86400

            # 基于文件新鲜度判断
            if days_since_modified <= 1:
                return '待处理'
            elif days_since_modified <= 7:
                return '处理中'
            elif os.path.getsize(file_path) > 0 and days_since_modified > 7:
                return '已完成'
            else:
                return '待处理'
        except Exception:
            return '有问题'
    
    def _extract_keywords(self, content: str, max_keywords: int = 10) -> List[str]:
        """提取关键词"""
        # 简单实现：基于词频统计
        # 实际应用中可以使用更复杂的NLP技术
        
        # 移除常见停用词
        stop_words = {'的', '了', '在', '是', '我', '有', '和', '就', '不', '人', '都', '一', '一个', '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看', '好', '自己', '这'}
        
        # 分词（简单按空格和标点分割）
        import re
        words = re.findall(r'[\u4e00-\u9fa5]+|[a-zA-Z]+', content)
        
        # 统计词频
        word_count = {}
        for word in words:
            if len(word) > 1 and word not in stop_words:
                word_count[word] = word_count.get(word, 0) + 1
        
        # 返回词频最高的关键词
        sorted_words = sorted(word_count.items(), key=lambda x: x[1], reverse=True)
        return [word for word, count in sorted_words[:max_keywords]]
    
    def _calculate_confidence(self, content: str, content_type: str) -> str:
        """计算分类置信度"""
        # 简单实现：基于内容长度和类型匹配度
        content_length = len(content)
        
        if content_length > 10000:
            return 'high'
        elif content_length > 1000:
            return 'medium'
        else:
            return 'low'
    
    def _get_classification_reason(self, content_type: str, quality: str) -> str:
        """获取分类原因"""
        reasons = {
            'API文档': '包含API接口相关关键词',
            '用户手册': '包含用户指南和操作说明',
            '技术文档': '包含技术实现和架构描述',
            '设计文档': '包含设计方案和规格说明',
            '会议纪要': '包含会议讨论和决议内容',
            '研究报告': '包含研究分析和数据报告',
            '培训材料': '包含培训教程和学习内容',
            '其他': '未匹配到特定类型特征'
        }
        
        return reasons.get(content_type, '基于内容分析')
    
    def _read_file_content(self, file_path: str) -> Optional[str]:
        """读取文件内容"""
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            # 对于文本文件，直接读取
            if file_ext in ['.txt', '.md', '.html', '.htm', '.rtf']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            
            # 对于PDF文件
            elif file_ext == '.pdf':
                try:
                    import pdfplumber
                    with pdfplumber.open(file_path) as pdf:
                        text_parts = []
                        for page in pdf.pages:
                            text = page.extract_text() or ""
                            text_parts.append(text)
                        return "\n".join(text_parts)
                except ImportError:
                    return None
            
            # 对于Word文档
            elif file_ext in ['.docx', '.doc']:
                try:
                    from docx import Document
                    doc = Document(file_path)
                    text_parts = []
                    for para in doc.paragraphs:
                        text_parts.append(para.text)
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                text_parts.append(cell.text)
                    return "\n".join(text_parts)
                except ImportError:
                    return None
            
            else:
                return None
                
        except Exception as e:
            logger.error(f"读取文件内容时出错: {file_path}, 错误: {e}")
            return None
    
    def _get_default_classification(self, file_path: str, reason: str) -> Dict[str, Any]:
        """获取默认分类结果"""
        return {
            'file_path': file_path,
            'file_name': os.path.basename(file_path),
            'content_type': '其他',
            'quality': '待评估',
            'processing_status': '有问题',
            'keywords': [],
            'classification_confidence': 'none',
            'classification_reason': reason
        }
    
    def batch_classify(self, file_paths: List[str]) -> Dict[str, Any]:
        """
        批量分类文档
        
        Args:
            file_paths: 文件路径列表
            
        Returns:
            批量分类结果
        """
        results = []
        
        for file_path in file_paths:
            result = self.classify_document(file_path)
            results.append(result)
        
        # 统计各类型数量
        content_type_counter = {}
        quality_counter = {}
        
        for result in results:
            content_type = result['content_type']
            quality = result['quality']
            
            content_type_counter[content_type] = content_type_counter.get(content_type, 0) + 1
            quality_counter[quality] = quality_counter.get(quality, 0) + 1
        
        summary = {
            'total_files': len(results),
            'content_type_distribution': content_type_counter,
            'quality_distribution': quality_counter,
            'classification_results': results
        }
        
        logger.info(f"批量文档分类完成: {len(results)} 个文件")
        return summary
    
    def get_classification_summary(self, classification_result: Dict[str, Any]) -> str:
        """获取分类结果摘要"""
        if not classification_result or classification_result['total_files'] == 0:
            return "没有文件需要分类"
        
        summary_lines = [
            f"文档分类结果摘要:",
            f"总文件数: {classification_result['total_files']}",
            "",
            "内容类型分布:"
        ]
        
        for content_type, count in classification_result['content_type_distribution'].items():
            percentage = round(count / classification_result['total_files'] * 100, 2)
            summary_lines.append(f"  {content_type}: {count} 个文件 ({percentage}%)")
        
        summary_lines.append("\n质量等级分布:")
        for quality, count in classification_result['quality_distribution'].items():
            percentage = round(count / classification_result['total_files'] * 100, 2)
            summary_lines.append(f"  {quality}: {count} 个文件 ({percentage}%)")
        
        return "\n".join(summary_lines)