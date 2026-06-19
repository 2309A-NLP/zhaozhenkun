"""
敏感信息检测模块
Sensitive Information Detection Module
"""

import os
import re
from typing import Dict, List, Any, Tuple, Optional
from .base import ConfigManager, logger


class SensitiveInfoDetector:
    """敏感信息检测器"""
    
    def __init__(self, config: ConfigManager):
        self.config = config
        self.enable_detection = config.get('sensitive_detection.enable_detection', True)
        self.detection_types = config.get('sensitive_detection.detection_types', {})
        self.context_chars = config.get('sensitive_detection.context_chars', 20)
        self.max_detection_per_file = config.get('sensitive_detection.max_detection_per_file', 50)
        self.max_pending_review = config.get('sensitive_detection.max_pending_review', 100)
        
        # 编译正则表达式
        self.patterns = {}
        for info_type, type_config in self.detection_types.items():
            if type_config.get('enabled', True):
                try:
                    self.patterns[info_type] = re.compile(type_config['pattern'])
                except re.error as e:
                    logger.error(f"编译正则表达式失败: {info_type}, 错误: {e}")
    
    def detect_sensitive_info(self, file_paths: List[str]) -> Dict[str, Any]:
        """
        检测敏感信息
        
        Args:
            file_paths: 文件路径列表
            
        Returns:
            敏感信息检测结果
        """
        if not file_paths or not self.enable_detection:
            logger.warning("没有文件需要检测或检测已禁用")
            return self._get_empty_result()
        
        all_detections = []
        pending_review = []
        files_with_sensitive = 0
        
        for file_path in file_paths:
            try:
                file_detections = self._detect_in_file(file_path)
                if file_detections:
                    files_with_sensitive += 1
                    all_detections.extend(file_detections)
                    
                    # 添加到待审核列表
                    for detection in file_detections:
                        pending_review.append({
                            'file_path': file_path,
                            'file_name': os.path.basename(file_path),
                            'info_type': detection['type'],
                            'value': detection['value'],
                            'context': detection['context'],
                            'line_number': detection.get('line_number'),
                            'position': detection.get('position')
                        })
            except Exception as e:
                logger.warning(f"检测文件敏感信息时出错: {file_path}, 错误: {e}")
                continue
        
        # 限制待审核列表数量
        if len(pending_review) > self.max_pending_review:
            pending_review = pending_review[:self.max_pending_review]
        
        # 统计各类型数量
        type_counter = {}
        for detection in all_detections:
            info_type = detection['type']
            type_counter[info_type] = type_counter.get(info_type, 0) + 1
        
        result = {
            'total_files': len(file_paths),
            'files_with_sensitive': files_with_sensitive,
            'total_detections': len(all_detections),
            'type_distribution': type_counter,
            'detections': all_detections,
            'pending_review_count': len(pending_review),
            'pending_review_list': pending_review
        }
        
        logger.info(f"敏感信息检测完成: {files_with_sensitive} 个文件包含敏感信息，共 {len(all_detections)} 处")
        return result
    
    def _detect_in_file(self, file_path: str) -> List[Dict[str, Any]]:
        """在单个文件中检测敏感信息"""
        try:
            content = self._read_file_content(file_path)
            if not content:
                return []
            
            detections = []
            lines = content.split('\n')
            
            for line_num, line in enumerate(lines, 1):
                for info_type, pattern in self.patterns.items():
                    matches = pattern.finditer(line)
                    for match in matches:
                        # 提取上下文
                        start = max(0, match.start() - self.context_chars)
                        end = min(len(line), match.end() + self.context_chars)
                        context = line[start:end]
                        
                        # 添加省略号表示截断
                        if start > 0:
                            context = "..." + context
                        if end < len(line):
                            context = context + "..."
                        
                        detection = {
                            'type': info_type,
                            'value': match.group(),
                            'context': context,
                            'line_number': line_num,
                            'position': match.start(),
                            'file_path': file_path
                        }
                        
                        detections.append(detection)
                        
                        # 限制每个文件的最大检测数量
                        if len(detections) >= self.max_detection_per_file:
                            logger.warning(f"文件 {file_path} 检测数量达到上限 {self.max_detection_per_file}")
                            return detections
            
            return detections
            
        except Exception as e:
            logger.error(f"读取文件内容时出错: {file_path}, 错误: {e}")
            return []
    
    def _read_file_content(self, file_path: str) -> Optional[str]:
        """读取文件内容"""
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            # 对于文本文件，直接读取
            if file_ext in ['.txt', '.md', '.html', '.htm', '.rtf']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            
            # 对于PDF文件
            elif file_ext == '.pdf':
                try:
                    import pdfplumber
                    with pdfplumber.open(file_path) as pdf:
                        text_parts = []
                        for page in pdf.pages:
                            text = page.extract_text() or ""
                            text_parts.append(text)
                        return "\n".join(text_parts)
                except ImportError:
                    logger.warning("pdfplumber未安装，无法处理PDF文件")
                    return None
            
            # 对于Word文档
            elif file_ext in ['.docx', '.doc']:
                try:
                    from docx import Document
                    doc = Document(file_path)
                    text_parts = []
                    for para in doc.paragraphs:
                        text_parts.append(para.text)
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                text_parts.append(cell.text)
                    return "\n".join(text_parts)
                except ImportError:
                    logger.warning("python-docx未安装，无法处理Word文档")
                    return None
            
            # 对于ODT文件
            elif file_ext == '.odt':
                try:
                    from odf import text, teletype
                    from odf.opendocument import load
                    doc = load(file_path)
                    text_parts = []
                    for para in doc.getElementsByType(text.P):
                        text_parts.append(teletype.extractText(para))
                    return "\n".join(text_parts)
                except ImportError:
                    logger.warning("odfpy未安装，无法处理ODT文件")
                    return None
            
            else:
                logger.warning(f"不支持的文件格式: {file_ext}")
                return None
                
        except Exception as e:
            logger.error(f"读取文件内容时出错: {file_path}, 错误: {e}")
            return None
    
    def _get_empty_result(self) -> Dict[str, Any]:
        """获取空结果"""
        return {
            'total_files': 0,
            'files_with_sensitive': 0,
            'total_detections': 0,
            'type_distribution': {},
            'detections': [],
            'pending_review_count': 0,
            'pending_review_list': []
        }
    
    def get_sensitive_summary(self, detection_result: Dict[str, Any]) -> str:
        """获取敏感信息检测摘要"""
        if not detection_result or detection_result['total_files'] == 0:
            return "没有文件需要检测"
        
        summary_lines = [
            f"敏感信息检测结果摘要:",
            f"总文件数: {detection_result['total_files']}",
            f"包含敏感信息的文件数: {detection_result['files_with_sensitive']}",
            f"敏感信息总数: {detection_result['total_detections']}",
            "",
            "各类型分布:"
        ]
        
        for info_type, count in detection_result['type_distribution'].items():
            summary_lines.append(f"  {info_type}: {count} 处")
        
        if detection_result['pending_review_count'] > 0:
            summary_lines.append(f"\n待审核项目: {detection_result['pending_review_count']} 项")
            summary_lines.append("（需要人工审核的敏感信息已添加到待审核列表）")
            
            # 显示部分待审核示例
            summary_lines.append("\n待审核示例:")
            for i, item in enumerate(detection_result['pending_review_list'][:3], 1):
                summary_lines.append(f"  {i}. 文件: {item['file_name']}")
                summary_lines.append(f"     类型: {item['info_type']}")
                summary_lines.append(f"     上下文: {item['context']}")
        
        return "\n".join(summary_lines)
    
    def validate_detection_patterns(self) -> Dict[str, bool]:
        """验证检测模式是否有效"""
        validation_results = {}
        
        for info_type, pattern in self.patterns.items():
            try:
                # 测试模式是否能匹配简单字符串
                test_string = "test123"
                result = pattern.search(test_string)
                validation_results[info_type] = True
            except Exception as e:
                logger.error(f"验证检测模式失败: {info_type}, 错误: {e}")
                validation_results[info_type] = False
        
        return validation_results