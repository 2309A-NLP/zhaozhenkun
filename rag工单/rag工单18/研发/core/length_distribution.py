"""
文档长度分布统计模块
Document Length Distribution Statistics Module
"""

import os
from typing import Dict, List, Any, Optional
from .base import ConfigManager, logger

try:
    import numpy as np
    NUMPY_AVAILABLE = True
except ImportError:
    NUMPY_AVAILABLE = False
    logger.warning("numpy未安装，分位数计算将使用标准库。请运行: pip install numpy")


class LengthDistributionAnalyzer:
    """文档长度分布分析器"""
    
    def __init__(self, config: ConfigManager):
        self.config = config
        self.length_bins = config.get('length_distribution.length_bins', [0, 100, 500, 1000, 5000, 10000, 50000, 100000, 500000])
        self.percentiles = config.get('length_distribution.percentiles', [25, 50, 75, 90, 99])
    
    def analyze(self, file_paths: List[str]) -> Dict[str, Any]:
        """
        分析文档长度分布
        
        Args:
            file_paths: 文件路径列表
            
        Returns:
            长度分布统计结果
        """
        if not file_paths:
            logger.warning("没有文件需要分析")
            return self._get_empty_result()
        
        # 收集所有文档的字符数
        char_counts = []
        file_lengths = []
        
        for file_path in file_paths:
            try:
                char_count = self._get_file_char_count(file_path)
                if char_count is not None:
                    char_counts.append(char_count)
                    file_lengths.append({
                        'file_path': file_path,
                        'file_name': os.path.basename(file_path),
                        'char_count': char_count
                    })
            except Exception as e:
                logger.warning(f"获取文件字符数时出错: {file_path}, 错误: {e}")
                continue
        
        if not char_counts:
            logger.warning("没有成功获取到文件字符数")
            return self._get_empty_result()
        
        # 计算统计指标
        statistics = self._calculate_statistics(char_counts)
        
        # 计算长度区间分布
        bin_distribution = self._calculate_bin_distribution(char_counts)
        
        # 按长度排序文件
        file_lengths.sort(key=lambda x: x['char_count'], reverse=True)
        
        result = {
            'total_files': len(char_counts),
            'statistics': statistics,
            'bin_distribution': bin_distribution,
            'longest_files': file_lengths[:10],  # 前10个最长的文件
            'shortest_files': file_lengths[-10:] if len(file_lengths) > 10 else []  # 后10个最短的文件
        }
        
        logger.info(f"文档长度分布分析完成: {len(char_counts)} 个文件")
        return result
    
    def _get_file_char_count(self, file_path: str) -> Optional[int]:
        """获取文件字符数"""
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            # 对于文本文件，直接读取字符数
            if file_ext in ['.txt', '.md', '.html', '.htm', '.rtf']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return len(f.read())
            
            # 对于PDF文件，使用pdfplumber
            elif file_ext == '.pdf':
                try:
                    import pdfplumber
                    with pdfplumber.open(file_path) as pdf:
                        total_chars = 0
                        for page in pdf.pages:
                            text = page.extract_text() or ""
                            total_chars += len(text)
                        return total_chars
                except ImportError:
                    logger.warning("pdfplumber未安装，无法处理PDF文件")
                    return None
            
            # 对于Word文档，使用python-docx
            elif file_ext in ['.docx', '.doc']:
                try:
                    from docx import Document
                    doc = Document(file_path)
                    total_chars = 0
                    for para in doc.paragraphs:
                        total_chars += len(para.text)
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                total_chars += len(cell.text)
                    return total_chars
                except ImportError:
                    logger.warning("python-docx未安装，无法处理Word文档")
                    return None
            
            # 对于ODT文件
            elif file_ext == '.odt':
                try:
                    from odf import text, teletype
                    from odf.opendocument import load
                    doc = load(file_path)
                    total_chars = 0
                    for para in doc.getElementsByType(text.P):
                        total_chars += len(teletype.extractText(para))
                    return total_chars
                except ImportError:
                    logger.warning("odfpy未安装，无法处理ODT文件")
                    return None
            
            else:
                logger.warning(f"不支持的文件格式: {file_ext}")
                return None
                
        except Exception as e:
            logger.error(f"读取文件字符数时出错: {file_path}, 错误: {e}")
            return None
    
    def _calculate_statistics(self, char_counts: List[int]) -> Dict[str, Any]:
        """计算统计指标"""
        if not char_counts:
            return {}
        
        # 基本统计
        total_chars = sum(char_counts)
        avg_chars = total_chars / len(char_counts)
        min_chars = min(char_counts)
        max_chars = max(char_counts)
        
        # 计算分位数
        percentiles = self._calculate_percentiles(char_counts)
        
        return {
            'total_chars': total_chars,
            'avg_chars': round(avg_chars, 2),
            'min_chars': min_chars,
            'max_chars': max_chars,
            'percentiles': percentiles
        }
    
    def _calculate_percentiles(self, char_counts: List[int]) -> Dict[str, float]:
        """计算分位数"""
        if NUMPY_AVAILABLE:
            # 使用numpy计算分位数
            percentile_values = np.percentile(char_counts, self.percentiles)
            return {f'P{p}': round(float(v), 2) for p, v in zip(self.percentiles, percentile_values)}
        else:
            # 使用标准库计算分位数
            sorted_counts = sorted(char_counts)
            n = len(sorted_counts)
            percentile_values = {}
            
            for p in self.percentiles:
                k = (n - 1) * p / 100
                f = int(k)
                c = f + 1 if f + 1 < n else f
                d = k - f
                value = sorted_counts[f] + d * (sorted_counts[c] - sorted_counts[f])
                percentile_values[f'P{p}'] = round(value, 2)
            
            return percentile_values
    
    def _calculate_bin_distribution(self, char_counts: List[int]) -> Dict[str, Any]:
        """计算长度区间分布"""
        # 初始化区间计数
        bin_counts = {}
        for i in range(len(self.length_bins) - 1):
            lower = self.length_bins[i]
            upper = self.length_bins[i + 1]
            bin_key = f"{lower}-{upper}"
            bin_counts[bin_key] = 0
        
        # 添加"以上"区间
        last_bin = f"{self.length_bins[-1]}以上"
        bin_counts[last_bin] = 0
        
        # 统计每个区间的文件数
        for count in char_counts:
            placed = False
            for i in range(len(self.length_bins) - 1):
                lower = self.length_bins[i]
                upper = self.length_bins[i + 1]
                if lower <= count < upper:
                    bin_key = f"{lower}-{upper}"
                    bin_counts[bin_key] += 1
                    placed = True
                    break
            
            if not placed and count >= self.length_bins[-1]:
                bin_counts[last_bin] += 1
        
        # 计算占比
        total_files = len(char_counts)
        bin_distribution = {}
        for bin_key, count in bin_counts.items():
            bin_distribution[bin_key] = {
                'count': count,
                'percentage': round(count / total_files * 100, 2) if total_files > 0 else 0
            }
        
        return bin_distribution
    
    def _get_empty_result(self) -> Dict[str, Any]:
        """获取空结果"""
        return {
            'total_files': 0,
            'statistics': {},
            'bin_distribution': {},
            'longest_files': [],
            'shortest_files': []
        }
    
    def get_length_summary(self, analysis_result: Dict[str, Any]) -> str:
        """获取长度分布摘要文本"""
        if not analysis_result or analysis_result['total_files'] == 0:
            return "没有文件需要分析"
        
        stats = analysis_result['statistics']
        summary_lines = [
            f"文档长度分布摘要:",
            f"总文件数: {analysis_result['total_files']}",
            "",
            "基本统计:",
            f"  总字符数: {stats.get('total_chars', 0):,}",
            f"  平均字符数: {stats.get('avg_chars', 0):,.2f}",
            f"  最小字符数: {stats.get('min_chars', 0):,}",
            f"  最大字符数: {stats.get('max_chars', 0):,}",
            "",
            "分位数分布:"
        ]
        
        percentiles = stats.get('percentiles', {})
        for p_name, p_value in percentiles.items():
            summary_lines.append(f"  {p_name}: {p_value:,.2f}")
        
        summary_lines.append("\n长度区间分布:")
        for bin_key, bin_info in analysis_result['bin_distribution'].items():
            summary_lines.append(f"  {bin_key}: {bin_info['count']} 个文件 ({bin_info['percentage']}%)")
        
        if analysis_result['longest_files']:
            summary_lines.append("\n最长的文件:")
            for i, file_info in enumerate(analysis_result['longest_files'][:5], 1):
                summary_lines.append(f"  {i}. {file_info['file_name']}: {file_info['char_count']:,} 字符")
        
        return "\n".join(summary_lines)